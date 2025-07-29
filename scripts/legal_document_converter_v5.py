#!/usr/bin/env python3
"""
Legal Document PDF to DOCX Converter v5
========================================

Enhanced converter with specialized handling for Versus sections and respondent lists
with proper line number alignment and formatting.
"""

import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class LegalDocumentConverterV5:
    """Enhanced legal document converter with Versus and respondent list handling."""
    
    def __init__(self):
        self.page_number_patterns = [
            r'\bpage\s*no\.?\s*\d+',
            r'\bpage\s*\d+',
            r'\d+\s*of\s*\d+',
            r'\(\d+\s*of\s*\d+\)',
            r'^\d+$',  # Standalone page numbers
            r'Page\s*\d+',
            r'PAGE\s*\d+'
        ]
        
        self.watermark_patterns = [
            r'CONFIDENTIAL',
            r'DRAFT',
            r'COPY',
            r'SCANNED',
            r'DIGITAL\s*COPY',
            r'ORIGINAL'
        ]
        
        # Legal document specific patterns
        self.legal_headers = [
            r'IN\s*THE\s*SUPREME\s*COURT\s*OF\s*INDIA',
            r'IN\s*THE\s*HIGH\s*COURT\s*OF\s*[A-Z\s]+',
            r'CRIMINAL\s*APPEAL\s*NO\.',
            r'CIVIL\s*APPEAL\s*NO\.',
            r'CRIMINAL\s*APPELLATE\s*JURISDICTION',
            r'CIVIL\s*APPELLATE\s*JURISDICTION'
        ]
        
        # Party name patterns
        self.party_patterns = [
            r'\.\.\.\s*APPELLANT\(S\)',
            r'\.\.\.\s*RESPONDENT\(S\)',
            r'VERSUS',
            r'VS\.',
            r'----\s*RESPONDENT',
            r'----\s*RESPONDENTS'
        ]
        
        # List detection patterns
        self.list_patterns = [
            r'^\d+\.\s*',  # 1. 2. 3.
            r'^\(\d+\)\s*',  # (1) (2) (3)
            r'^[a-z]\)\s*',  # a) b) c)
            r'^[A-Z]\)\s*',  # A) B) C)
            r'^[ivx]+\.\s*',  # i. ii. iii. iv. v. vi. vii. viii. ix. x.
            r'^[IVX]+\.\s*',  # I. II. III. IV. V. VI. VII. VIII. IX. X.
        ]
        
        # Versus and respondent patterns
        self.versus_patterns = [
            r'^VERSUS$',
            r'^VS\.$',
            r'^Versus$',
            r'^vs\.$'
        ]
        
        self.respondent_patterns = [
            r'----\s*RESPONDENT',
            r'----\s*RESPONDENTS',
            r'\.\.\.\s*RESPONDENT\(S\)',
            r'RESPONDENT\(S\)'
        ]

    def is_page_number(self, text):
        """Check if text is a page number."""
        if not text or len(text.strip()) < 1:
            return False
        
        text_clean = text.strip()
        
        for pattern in self.page_number_patterns:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True
        
        return False

    def is_watermark(self, text):
        """Check if text is a watermark."""
        if not text:
            return False
        
        text_upper = text.upper().strip()
        
        for pattern in self.watermark_patterns:
            if re.search(pattern, text_upper):
                return True
        
        return False

    def clean_text(self, text):
        """Clean text while preserving legal document formatting."""
        if not text:
            return ""
        
        # Remove control characters except tabs and newlines
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize whitespace but preserve intentional spacing
        text = re.sub(r'[ \t]+', ' ', text)
        
        return text.strip()

    def detect_alignment(self, block, page_width=595):
        """Detect text alignment based on block position."""
        try:
            block_width = block[2] - block[0]  # Block width
            block_center = (block[0] + block[2]) / 2
            page_center = page_width / 2
            
            # If block spans most of the page width, it's likely justified
            if block_width > page_width * 0.8:
                return WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Check if text is centered (common for titles and headers)
            if abs(block_center - page_center) < page_width * 0.15:
                return WD_ALIGN_PARAGRAPH.CENTER
            
            # Check if text is right-aligned (common for dates, case numbers)
            if block_center > page_center + page_width * 0.15:
                return WD_ALIGN_PARAGRAPH.RIGHT
            
            # Default to left alignment
            return WD_ALIGN_PARAGRAPH.LEFT
        except:
            return WD_ALIGN_PARAGRAPH.LEFT

    def detect_font_size(self, text, block_position):
        """Detect appropriate font size based on content and position."""
        # Check if it's a main title (at top of page, all caps)
        if text.isupper() and len(text.strip()) < 100 and block_position[1] < 150:
            return Pt(16)
        
        # Check if it's a legal header
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return Pt(14)
        
        # Check if it's a case number or reference
        if re.search(r'(NO\.|CASE|APPEAL|PETITION)', text, re.IGNORECASE):
            return Pt(12)
        
        # Check if it's a party name (usually in all caps)
        if text.isupper() and len(text.strip()) < 50:
            return Pt(12)
        
        # Default to body text
        return Pt(11)

    def should_bold(self, text):
        """Determine if text should be bold."""
        # Titles and headings
        if text.isupper() and len(text.strip()) < 100:
            return True
        
        # Legal headers
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Case numbers and references
        if re.search(r'(NO\.|CASE|APPEAL|PETITION)', text, re.IGNORECASE):
            return True
        
        # Party names
        if text.isupper() and len(text.strip()) < 50:
            return True
        
        # Document type indicators
        if text.strip().upper() in ['JUDGMENT', 'ORDER', 'REPORTABLE']:
            return True
        
        return False

    def should_underline(self, text):
        """Determine if text should be underlined."""
        # Document type indicators
        if text.strip().upper() in ['JUDGMENT', 'ORDER', 'REPORTABLE']:
            return True
        
        # Case numbers with underscores
        if re.search(r'NO\.\s*_+\s*OF\s*\d+', text, re.IGNORECASE):
            return True
        
        # Specific legal patterns
        if re.search(r'CRIMINAL\s*APPEAL\s*NO\.', text, re.IGNORECASE):
            return True
        
        return False

    def is_legal_header(self, text):
        """Check if text is a legal document header."""
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False

    def is_party_name(self, text):
        """Check if text is a party name."""
        # Check for party patterns
        for pattern in self.party_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Check for all caps names (typical for parties)
        if text.isupper() and len(text.strip()) < 50 and not self.is_legal_header(text):
            return True
        
        return False

    def is_versus(self, text):
        """Check if text is 'Versus'."""
        for pattern in self.versus_patterns:
            if re.match(pattern, text.strip()):
                return True
        return False

    def is_respondent_section(self, text):
        """Check if text indicates respondent section."""
        for pattern in self.respondent_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False

    def is_list_item(self, text):
        """Check if text is a list item."""
        for pattern in self.list_patterns:
            if re.match(pattern, text.strip()):
                return True
        return False

    def is_list_continuation(self, text, prev_text):
        """Check if text is a continuation of a list item."""
        # If previous text ends with a list pattern and current text doesn't start with one
        if prev_text and not self.is_list_item(text):
            # Check if previous text ends with list pattern
            for pattern in self.list_patterns:
                if re.search(pattern + r'[^.]*$', prev_text.strip()):
                    return True
        return False

    def extract_list_number(self, text):
        """Extract list number from text."""
        for pattern in self.list_patterns:
            match = re.match(pattern, text.strip())
            if match:
                return match.group(0).strip()
        return None

    def is_respondent_list_item(self, text):
        """Check if text is a respondent list item."""
        # Check for patterns like "1. State of Rajasthan" or "1.State of Rajasthan"
        patterns = [
            r'^\d+\.\s*[A-Z]',  # 1. State
            r'^\d+[A-Z]',       # 1State (no space)
            r'^\d+\.\s*[A-Z][a-z]+\s+of\s+[A-Z]',  # 1. State of Rajasthan
        ]
        
        for pattern in patterns:
            if re.match(pattern, text.strip()):
                return True
        return False

    def merge_versus_and_respondents(self, blocks):
        """Merge Versus section and respondent lists properly."""
        merged_blocks = []
        i = 0
        
        while i < len(blocks):
            current_block = blocks[i]
            current_text = current_block['text']
            
            # Check if this is "Versus"
            if self.is_versus(current_text):
                # Add Versus as a separate centered block
                versus_block = current_block.copy()
                versus_block['is_versus'] = True
                versus_block['alignment'] = WD_ALIGN_PARAGRAPH.CENTER
                merged_blocks.append(versus_block)
                i += 1
                continue
            
            # Check if this is a respondent list item
            if self.is_respondent_list_item(current_text):
                # Look ahead to see if next blocks are continuations
                merged_text = current_text
                j = i + 1
                
                while j < len(blocks):
                    next_block = blocks[j]
                    next_text = next_block['text']
                    
                    # If next block is continuation of respondent list item
                    if (not self.is_list_item(next_text) and 
                        not self.is_versus(next_text) and 
                        not self.is_respondent_section(next_text) and
                        not self.is_legal_header(next_text)):
                        
                        # Check if it's a continuation (not starting with number)
                        if not re.match(r'^\d+', next_text.strip()):
                            merged_text += " " + next_text
                            j += 1
                            continue
                    
                    break
                
                # Create merged respondent block
                respondent_block = current_block.copy()
                respondent_block['text'] = merged_text
                respondent_block['is_respondent_item'] = True
                respondent_block['list_number'] = self.extract_list_number(current_text)
                
                merged_blocks.append(respondent_block)
                i = j  # Skip processed blocks
                continue
            
            # Check if this is respondent section indicator
            if self.is_respondent_section(current_text):
                # Add as separate block
                respondent_section_block = current_block.copy()
                respondent_section_block['is_respondent_section'] = True
                merged_blocks.append(respondent_section_block)
                i += 1
                continue
            
            # Check if this is a regular list item
            if self.is_list_item(current_text):
                # Look ahead to see if next block is continuation
                if i + 1 < len(blocks):
                    next_block = blocks[i + 1]
                    next_text = next_block['text']
                    
                    # If next block is continuation, merge them
                    if self.is_list_continuation(next_text, current_text):
                        # Merge the blocks
                        merged_text = current_text + " " + next_text
                        merged_block = current_block.copy()
                        merged_block['text'] = merged_text
                        merged_block['is_list_item'] = True
                        merged_block['list_number'] = self.extract_list_number(current_text)
                        
                        merged_blocks.append(merged_block)
                        i += 2  # Skip next block since we merged it
                        continue
            
            # If not a special case, add as is
            current_block['is_list_item'] = self.is_list_item(current_text)
            if current_block['is_list_item']:
                current_block['list_number'] = self.extract_list_number(current_text)
            
            merged_blocks.append(current_block)
            i += 1
        
        return merged_blocks

    def analyze_document_structure(self, doc):
        """Analyze the document to identify headers and footers."""
        logger.info("Analyzing document structure...")
        
        all_blocks = []
        top_blocks = []
        bottom_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            
            if blocks:
                # Sort blocks by vertical position
                blocks.sort(key=lambda b: b[1])
                
                # Collect top and bottom blocks
                if blocks:
                    top_blocks.append(blocks[0][4].strip())
                    bottom_blocks.append(blocks[-1][4].strip())
                
                all_blocks.extend(blocks)
        
        # Find frequent headers and footers
        top_counter = Counter(top_blocks)
        bottom_counter = Counter(bottom_blocks)
        
        # Identify headers that appear on most pages
        threshold = max(2, int(0.6 * len(doc)))
        frequent_headers = {text for text, count in top_counter.items() 
                          if count >= threshold and text.strip()}
        frequent_footers = {text for text, count in bottom_counter.items() 
                           if count >= threshold and text.strip()}
        
        logger.info(f"Found {len(frequent_headers)} frequent headers and {len(frequent_footers)} frequent footers")
        
        return frequent_headers, frequent_footers

    def process_legal_document_blocks(self, blocks, page_width=595):
        """Process blocks to maintain legal document structure."""
        processed_blocks = []
        
        for block in blocks:
            text = block[4].strip()
            if not text:
                continue
            
            # Clean the text
            cleaned_text = self.clean_text(text)
            if not cleaned_text:
                continue
            
            # Create processed block with metadata
            processed_block = {
                'text': cleaned_text,
                'block': block,
                'alignment': self.detect_alignment(block, page_width),
                'font_size': self.detect_font_size(cleaned_text, block),
                'bold': self.should_bold(cleaned_text),
                'underline': self.should_underline(cleaned_text),
                'is_legal_header': self.is_legal_header(cleaned_text),
                'is_party_name': self.is_party_name(cleaned_text),
                'is_case_number': bool(re.search(r'(NO\.|CASE|APPEAL|PETITION)', cleaned_text, re.IGNORECASE)),
                'position': (block[1], block[0])  # (y, x) for sorting
            }
            
            processed_blocks.append(processed_block)
        
        return processed_blocks

    def create_legal_document_structure(self, document, blocks):
        """Create proper legal document structure with correct formatting."""
        if not blocks:
            return
        
        # Sort blocks by vertical position, then horizontal
        blocks.sort(key=lambda b: b['position'])
        
        # Merge Versus and respondent sections
        merged_blocks = self.merge_versus_and_respondents(blocks)
        
        current_paragraph = None
        current_formatting = None
        
        for i, block in enumerate(merged_blocks):
            text = block['text']
            
            # Handle Versus specially
            if block.get('is_versus', False):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.font.size = Pt(12)
                run.bold = True
                run.font.name = 'Times New Roman'
                document.add_paragraph()  # Add blank line after Versus
                continue
            
            # Handle respondent section indicator
            if block.get('is_respondent_section', False):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(text)
                run.font.size = Pt(11)
                run.bold = True
                run.font.name = 'Times New Roman'
                continue
            
            # Handle respondent list items
            if block.get('is_respondent_item', False):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Use numbered list style for respondent items
                paragraph.style = document.styles['List Number']
                
                # Remove the number from the text (since style will add it)
                list_number = block.get('list_number', '')
                if list_number:
                    text_without_number = re.sub(r'^' + re.escape(list_number) + r'\s*', '', text)
                    run = paragraph.add_run(text_without_number)
                else:
                    run = paragraph.add_run(text)
                
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                continue
            
            # Determine if this should be a new paragraph
            should_new_paragraph = True
            
            # Check if this block should be part of the previous paragraph
            if i > 0 and current_formatting:
                prev_block = merged_blocks[i-1]
                
                # Same formatting and small gap = same paragraph
                if (block['alignment'] == current_formatting['alignment'] and
                    block['font_size'] == current_formatting['font_size'] and
                    block['bold'] == current_formatting['bold'] and
                    block['underline'] == current_formatting['underline']):
                    
                    # Check vertical gap
                    vertical_gap = block['block'][1] - (prev_block['block'][1] + prev_block['block'][3])
                    if vertical_gap < 25:  # Small gap
                        should_new_paragraph = False
            
            if should_new_paragraph:
                # Create new paragraph
                paragraph = document.add_paragraph()
                paragraph.alignment = block['alignment']
                
                # Handle list items specially
                if block.get('is_list_item', False):
                    # Use numbered list style for list items
                    paragraph.style = document.styles['List Number']
                    # Add the text without the number (since style will add it)
                    list_number = block.get('list_number', '')
                    if list_number:
                        # Remove the number from the text
                        text_without_number = re.sub(r'^' + re.escape(list_number) + r'\s*', '', text)
                        run = paragraph.add_run(text_without_number)
                    else:
                        run = paragraph.add_run(text)
                else:
                    # Regular paragraph
                    run = paragraph.add_run(text)
                
                run.font.size = block['font_size']
                run.bold = block['bold']
                if block['underline']:
                    run.underline = True
                
                # Set font to Times New Roman for legal documents
                run.font.name = 'Times New Roman'
                
                current_paragraph = paragraph
                current_formatting = {
                    'alignment': block['alignment'],
                    'font_size': block['font_size'],
                    'bold': block['bold'],
                    'underline': block['underline']
                }
            else:
                # Add to existing paragraph
                if current_paragraph:
                    # Add space before new text
                    current_paragraph.add_run(" ")
                    run = current_paragraph.add_run(text)
                    run.font.size = block['font_size']
                    run.bold = block['bold']
                    if block['underline']:
                        run.underline = True
                    run.font.name = 'Times New Roman'
            
            # Add extra spacing for legal headers and section breaks
            if block['is_legal_header'] or block['is_case_number']:
                document.add_paragraph()  # Add blank line

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX with exact legal document formatting preservation."""
        logger.info(f"Starting legal document conversion: {pdf_path} -> {docx_path}")
        
        try:
            # Open PDF
            doc = fitz.open(pdf_path)
            if not doc:
                raise ValueError("Failed to open PDF file")
            
            # Create Word document
            document = Document()
            
            # Set document margins for legal documents
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Analyze document structure
            frequent_headers, frequent_footers = self.analyze_document_structure(doc)
            
            # Process each page
            for page_num in range(len(doc)):
                logger.info(f"Processing page {page_num + 1}/{len(doc)}")
                
                page = doc[page_num]
                blocks = page.get_text("blocks")
                
                if not blocks:
                    continue
                
                # Get page dimensions
                page_width = page.rect.width
                
                # Filter out headers, footers, page numbers, and artifacts
                filtered_blocks = []
                for block in blocks:
                    text = block[4].strip()
                    
                    # Skip empty blocks
                    if not text:
                        continue
                    
                    # Skip frequent headers and footers
                    if text in frequent_headers or text in frequent_footers:
                        logger.debug(f"Skipping frequent header/footer: {text[:50]}...")
                        continue
                    
                    # Skip page numbers
                    if self.is_page_number(text):
                        logger.debug(f"Skipping page number: {text}")
                        continue
                    
                    # Skip watermarks
                    if self.is_watermark(text):
                        logger.debug(f"Skipping watermark: {text}")
                        continue
                    
                    filtered_blocks.append(block)
                
                # Process blocks for legal document structure
                processed_blocks = self.process_legal_document_blocks(filtered_blocks, page_width)
                
                # Create document structure
                self.create_legal_document_structure(document, processed_blocks)
                
                # Add page break between pages (except for the last page)
                if page_num < len(doc) - 1:
                    document.add_page_break()
            
            # Save the document
            document.save(docx_path)
            logger.info(f"Successfully converted legal document to DOCX: {docx_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error during conversion: {str(e)}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()

def main():
    """Main function for command-line usage."""
    if len(sys.argv) != 3:
        print("Usage: python legal_document_converter_v5.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    # Validate input file
    if not os.path.exists(pdf_path):
        print(f"Error: Input PDF file not found: {pdf_path}", file=sys.stderr)
        sys.exit(1)
    
    # Create output directory if needed
    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert PDF to DOCX
    converter = LegalDocumentConverterV5()
    try:
        success = converter.convert_pdf_to_docx(pdf_path, docx_path)
        if success:
            print(f"Successfully converted {pdf_path} to {docx_path}")
        else:
            print("Conversion failed", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main() 