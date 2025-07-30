#!/usr/bin/env python3
"""
DOCX Document Analyzer
=====================

Comprehensive tool to analyze DOCX documents and identify formatting issues,
particularly focusing on numbered lists, alignment, and document structure.
"""

import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import defaultdict, Counter
import json
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocxAnalyzer:
    """Comprehensive DOCX document analyzer for formatting issues."""
    
    def __init__(self):
        self.numbered_list_patterns = [
            r'^\d+\.\s+',  # 1. text
            r'^\d+\)\s+',  # 1) text
            r'^[a-z]\)\s+',  # a) text
            r'^[A-Z]\)\s+',  # A) text
            r'^[ivxlcdm]+\.\s+',  # i. ii. iii. etc.
            r'^[IVXLCDM]+\.\s+',  # I. II. III. etc.
        ]
        
        self.alignment_issues = []
        self.numbered_list_issues = []
        self.formatting_issues = []
        self.structure_issues = []
        
    def analyze_document(self, docx_path):
        """Analyze a DOCX document for formatting issues."""
        logger.info(f"Analyzing document: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            # Reset analysis results
            self.alignment_issues = []
            self.numbered_list_issues = []
            self.formatting_issues = []
            self.structure_issues = []
            
            # Analyze paragraphs
            self.analyze_paragraphs(doc)
            
            # Analyze document structure
            self.analyze_document_structure(doc)
            
            # Generate comprehensive report
            report = self.generate_analysis_report(docx_path)
            
            return report
            
        except Exception as e:
            logger.error(f"Error analyzing document: {str(e)}")
            raise
    
    def analyze_paragraphs(self, doc):
        """Analyze individual paragraphs for formatting issues."""
        logger.info("Analyzing paragraphs...")
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check for numbered list issues
            self.check_numbered_list_formatting(para, i, text)
            
            # Check alignment issues
            self.check_alignment_issues(para, i, text)
            
            # Check formatting consistency
            self.check_formatting_consistency(para, i, text)
    
    def check_numbered_list_formatting(self, para, para_index, text):
        """Check for numbered list formatting issues."""
        # Check if this looks like a numbered list item
        is_numbered = False
        list_pattern = None
        
        for pattern in self.numbered_list_patterns:
            if re.match(pattern, text):
                is_numbered = True
                list_pattern = pattern
                break
        
        if is_numbered:
            # Check if the number is properly aligned with the text
            runs = para.runs
            if len(runs) > 0:
                first_run = runs[0]
                
                # Check if the number and text are in the same run
                if len(runs) == 1:
                    # Single run - check if there's proper spacing after the number
                    match = re.match(list_pattern, text)
                    if match:
                        number_part = match.group(0)
                        text_part = text[len(number_part):]
                        
                        if not text_part.startswith(' '):
                            self.numbered_list_issues.append({
                                'paragraph_index': para_index,
                                'issue': 'Missing space after number',
                                'text': text[:100],
                                'suggestion': 'Add space between number and text'
                            })
                
                # Check alignment for numbered lists
                if para.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    self.numbered_list_issues.append({
                        'paragraph_index': para_index,
                        'issue': 'Numbered list should be left-aligned',
                        'text': text[:100],
                        'current_alignment': str(para.alignment),
                        'suggestion': 'Set alignment to LEFT'
                    })
    
    def check_alignment_issues(self, para, para_index, text):
        """Check for alignment issues."""
        # Check for inconsistent alignment patterns
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            # Center-aligned text should typically be short (titles, headers)
            if len(text) > 200:
                self.alignment_issues.append({
                    'paragraph_index': para_index,
                    'issue': 'Long text center-aligned',
                    'text': text[:100],
                    'suggestion': 'Consider left alignment for long text'
                })
        
        # Check for justified text that might be too short
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            if len(text) < 50:
                self.alignment_issues.append({
                    'paragraph_index': para_index,
                    'issue': 'Short text justified',
                    'text': text[:100],
                    'suggestion': 'Consider left alignment for short text'
                })
    
    def check_formatting_consistency(self, para, para_index, text):
        """Check for formatting consistency issues."""
        runs = para.runs
        
        if len(runs) > 1:
            # Check for inconsistent font sizes
            font_sizes = [run.font.size for run in runs if run.font.size]
            if font_sizes and len(set(font_sizes)) > 1:
                self.formatting_issues.append({
                    'paragraph_index': para_index,
                    'issue': 'Inconsistent font sizes in paragraph',
                    'text': text[:100],
                    'font_sizes': [str(size) for size in font_sizes],
                    'suggestion': 'Use consistent font size within paragraph'
                })
            
            # Check for inconsistent bold formatting
            bold_states = [run.bold for run in runs]
            if len(set(bold_states)) > 1:
                self.formatting_issues.append({
                    'paragraph_index': para_index,
                    'issue': 'Inconsistent bold formatting in paragraph',
                    'text': text[:100],
                    'suggestion': 'Use consistent bold formatting within paragraph'
                })
    
    def analyze_document_structure(self, doc):
        """Analyze overall document structure."""
        logger.info("Analyzing document structure...")
        
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        
        # Check for document hierarchy
        heading_patterns = [
            r'^IN\s+THE\s+SUPREME\s+COURT',
            r'^IN\s+THE\s+HIGH\s+COURT',
            r'^CRIMINAL\s+APPEAL',
            r'^CIVIL\s+APPEAL',
            r'^JUDGMENT$',
            r'^ORDER$',
            r'^REPORTABLE$'
        ]
        
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            
            # Check if this should be a heading
            for pattern in heading_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    # Check if it's properly formatted as a heading
                    if not para.runs or not para.runs[0].bold:
                        self.structure_issues.append({
                            'paragraph_index': i,
                            'issue': 'Legal header not bold',
                            'text': text,
                            'suggestion': 'Make legal headers bold'
                        })
                    
                    # Check if it's properly centered
                    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        self.structure_issues.append({
                            'paragraph_index': i,
                            'issue': 'Legal header not centered',
                            'text': text,
                            'suggestion': 'Center legal headers'
                        })
                    break
        
        # Check for versus section formatting
        versus_patterns = [r'VERSUS', r'VS\.', r'v\.']
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if any(re.search(pattern, text, re.IGNORECASE) for pattern in versus_patterns):
                if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.structure_issues.append({
                        'paragraph_index': i,
                        'issue': 'Versus section not centered',
                        'text': text,
                        'suggestion': 'Center versus sections'
                    })
    
    def generate_analysis_report(self, docx_path):
        """Generate a comprehensive analysis report."""
        logger.info("Generating analysis report...")
        
        report = {
            'document_path': docx_path,
            'analysis_summary': {
                'total_issues': len(self.alignment_issues) + len(self.numbered_list_issues) + 
                               len(self.formatting_issues) + len(self.structure_issues),
                'alignment_issues': len(self.alignment_issues),
                'numbered_list_issues': len(self.numbered_list_issues),
                'formatting_issues': len(self.formatting_issues),
                'structure_issues': len(self.structure_issues)
            },
            'detailed_issues': {
                'alignment_issues': self.alignment_issues,
                'numbered_list_issues': self.numbered_list_issues,
                'formatting_issues': self.formatting_issues,
                'structure_issues': self.structure_issues
            },
            'recommendations': self.generate_recommendations()
        }
        
        return report
    
    def generate_recommendations(self):
        """Generate improvement recommendations."""
        recommendations = []
        
        if self.numbered_list_issues:
            recommendations.append({
                'category': 'Numbered Lists',
                'priority': 'High',
                'description': 'Improve numbered list formatting',
                'actions': [
                    'Ensure proper spacing after numbers',
                    'Use consistent left alignment for lists',
                    'Maintain proper indentation'
                ]
            })
        
        if self.alignment_issues:
            recommendations.append({
                'category': 'Alignment',
                'priority': 'Medium',
                'description': 'Fix alignment inconsistencies',
                'actions': [
                    'Use left alignment for body text',
                    'Center short titles and headers',
                    'Avoid justifying short text'
                ]
            })
        
        if self.structure_issues:
            recommendations.append({
                'category': 'Document Structure',
                'priority': 'High',
                'description': 'Improve legal document structure',
                'actions': [
                    'Make legal headers bold and centered',
                    'Center versus sections',
                    'Maintain proper document hierarchy'
                ]
            })
        
        return recommendations
    
    def print_report(self, report):
        """Print a formatted analysis report."""
        print("\n" + "="*80)
        print("DOCX DOCUMENT ANALYSIS REPORT")
        print("="*80)
        print(f"Document: {report['document_path']}")
        print()
        
        summary = report['analysis_summary']
        print("SUMMARY:")
        print(f"  Total Issues: {summary['total_issues']}")
        print(f"  Alignment Issues: {summary['alignment_issues']}")
        print(f"  Numbered List Issues: {summary['numbered_list_issues']}")
        print(f"  Formatting Issues: {summary['formatting_issues']}")
        print(f"  Structure Issues: {summary['structure_issues']}")
        print()
        
        if summary['total_issues'] == 0:
            print("✅ No formatting issues found! Document appears to be well-formatted.")
        else:
            print("ISSUES FOUND:")
            print()
            
            # Print numbered list issues
            if report['detailed_issues']['numbered_list_issues']:
                print("🔢 NUMBERED LIST ISSUES:")
                for issue in report['detailed_issues']['numbered_list_issues']:
                    print(f"  Paragraph {issue['paragraph_index']}: {issue['issue']}")
                    print(f"    Text: {issue['text']}")
                    print(f"    Suggestion: {issue['suggestion']}")
                    print()
            
            # Print alignment issues
            if report['detailed_issues']['alignment_issues']:
                print("📐 ALIGNMENT ISSUES:")
                for issue in report['detailed_issues']['alignment_issues']:
                    print(f"  Paragraph {issue['paragraph_index']}: {issue['issue']}")
                    print(f"    Text: {issue['text']}")
                    print(f"    Suggestion: {issue['suggestion']}")
                    print()
            
            # Print structure issues
            if report['detailed_issues']['structure_issues']:
                print("🏗️  STRUCTURE ISSUES:")
                for issue in report['detailed_issues']['structure_issues']:
                    print(f"  Paragraph {issue['paragraph_index']}: {issue['issue']}")
                    print(f"    Text: {issue['text']}")
                    print(f"    Suggestion: {issue['suggestion']}")
                    print()
            
            # Print recommendations
            if report['recommendations']:
                print("💡 RECOMMENDATIONS:")
                for rec in report['recommendations']:
                    print(f"  {rec['category']} ({rec['priority']} priority):")
                    print(f"    {rec['description']}")
                    for action in rec['actions']:
                        print(f"    • {action}")
                    print()
        
        print("="*80)

def main():
    """Main function for command-line usage."""
    if len(sys.argv) != 2:
        print("Usage: python docx_analyzer.py <document.docx>", file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    # Validate input file
    if not os.path.exists(docx_path):
        print(f"Error: DOCX file not found: {docx_path}", file=sys.stderr)
        sys.exit(1)
    
    # Analyze document
    analyzer = DocxAnalyzer()
    try:
        report = analyzer.analyze_document(docx_path)
        analyzer.print_report(report)
        
        # Save detailed report to JSON file
        report_path = docx_path.replace('.docx', '_analysis_report.json')
        with open(report_path, 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
        
        print(f"\nDetailed report saved to: {report_path}")
        
    except Exception as e:
        print(f"Error during analysis: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main() 