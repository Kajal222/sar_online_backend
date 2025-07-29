from docx import Document
import json
import sys
import os

def generate_docx(meta, paragraphs, output_path):
    doc = Document()
    doc.add_heading("Judgment Metadata", level=1)
    for key, val in meta.items():
        doc.add_paragraph(f"{key.title()}: {val}")
    doc.add_heading("Judgment Text", level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)

if __name__ == "__main__":
    def read_json_arg(arg):
        try:
            if arg == '-':
                return json.loads(sys.stdin.readline())
            else:
                if not os.path.exists(arg):
                    print(f"File not found: {arg}", file=sys.stderr)
                    sys.exit(1)
                with open(arg, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            print(f"Failed to read or parse JSON from {arg}: {e}", file=sys.stderr)
            sys.exit(1)
    if len(sys.argv) < 4:
        print("Usage: python docx_generator.py <meta.json> <paragraphs.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    meta = read_json_arg(sys.argv[1])
    paragraphs = read_json_arg(sys.argv[2])
    output_path = sys.argv[3]
    generate_docx(meta, paragraphs, output_path)
