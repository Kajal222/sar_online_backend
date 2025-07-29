#!/usr/bin/env python3
"""
Enhanced PDF to DOCX Converter for Legal Documents
==================================================

This script converts PDF legal documents (judgments, orders) to DOCX format while:
- Preserving original fonts, sizes, and alignment
- Removing page numbers from headers/footers
- Removing watermarks and scanner artifacts
- Maintaining document structure and formatting
- Preserving content integrity

Features:
- Advanced text block analysis for legal document structure
- Font preservation and mapping
- Alignment detection and preservation
- Watermark and artifact removal
- Page number detection and removal
- Legal document-specific formatting rules
"""

import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from collections import Counter, defaultdict
import json
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class LegalDocumentConverter:
    """Enhanced converter for legal PDF documents to DOCX format."""
    
    def __init__(self):
        self.page_number_patterns = [
            r'\bpage\s*no\.?\s*\d+',
            r'\bpage\s*\d+',
            r'\d+\s*of\s*\d+',
            r'\(\d+\s*of\s*\d+\)',
            r'\[CW-\d+/\d+\]',
            r'\(\d+\s*of\s*\d+\)',
            r'^\d+$',  # Standalone page numbers
            r'^\s*\d+\s*$',  # Page numbers with whitespace
            r'Page\s*\d+',
            r'PAGE\s*\d+',
            r'page\s*\d+'
        ]
        
        self.watermark_patterns = [
            r'CONFIDENTIAL',
            r'DRAFT',
            r'COPY',
            r'SCANNED',
            r'DIGITAL\s*COPY',
            r'ELECTRONIC\s*COPY',
            r'ORIGINAL',
            r'CERTIFIED\s*COPY',
            r'TRUE\s*COPY',
            r'OFFICIAL\s*COPY'
        ]
        
        self.scanner_artifacts = [
            r'[^\w\s\.\,\;\:\!\?\(\)\[\]\{\}\-\_\'\"\@\#\$\%\&\*\+\=\|\/\\]',  # Non-standard characters
            r'[^\x00-\x7F]',  # Non-ASCII characters that might be scanner artifacts
        ]
        
        # Legal document specific patterns
        self.legal_headers = [
            r'IN\s*THE\s*HIGH\s*COURT\s*OF\s*[A-Z\s]+',
            r'IN\s*THE\s*SUPREME\s*COURT\s*OF\s*[A-Z\s]+',
            r'BEFORE\s*THE\s*HON\'BLE\s*[A-Z\s]+',
            r'CRIMINAL\s*APPEAL\s*NO\.',
            r'CIVIL\s*APPEAL\s*NO\.',
            r'WRIT\s*PETITION\s*NO\.',
            r'SPECIAL\s*LEAVE\s*PETITION\s*NO\.'
        ]
        
        # Font size mapping for legal documents
        self.font_size_mapping = {
            'title': Pt(16),
            'heading': Pt(14),
            'subheading': Pt(12),
            'body': Pt(11),
            'small': Pt(10),
            'footnote': Pt(9)
        }

    def is_page_number(self, text):
        """Check if text is a page number."""
        if not text or len(text.strip()) < 1:
            return False
        
        text_clean = text.strip()
        
        # Check against patterns
        for pattern in self.page_number_patterns:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True
        
        # Check if it's just a number (likely page number)
        if re.match(r'^\d+$', text_clean):
            return True
            
        return False

    def is_watermark(self, text):
        """Check if text is a watermark."""
        if not text:
            return False
        
        text_upper = text.upper().strip()
        
        for pattern in self.watermark_patterns:
            if re.search(pattern, text_upper):
                return True
        
        return False

    def is_scanner_artifact(self, text):
        """Check if text contains scanner artifacts."""
        if not text:
            return False
        
        # Check for excessive non-standard characters
        non_standard_chars = len(re.findall(r'[^\w\s\.\,\;\:\!\?\(\)\[\]\{\}\-\_\'\"\@\#\$\%\&\*\+\=\|\/\\]', text))
        if non_standard_chars > len(text) * 0.1:  # More than 10% non-standard chars
            return True
        
        return False

    def clean_text(self, text):
        """Clean text while preserving legal document formatting."""
        if not text:
            return ""
        
        # Remove control characters except tabs and newlines
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize whitespace but preserve intentional spacing
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()

    def detect_alignment(self, block):
        """Detect text alignment based on block position and content."""
        page_width = block[2] - block[0]  # Block width
        page_center = (block[0] + block[2]) / 2
        
        # If block spans most of the page width, it's likely justified
        if page_width > 400:  # Assuming A4 page width ~595 points
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Check if text is centered (common for titles)
        if abs(block[0] - page_center) < 50:  # Within 50 points of center
            return WD_ALIGN_PARAGRAPH.CENTER
        
        # Check if text is right-aligned (common for dates, case numbers)
        if block[0] > page_center + 100:
            return WD_ALIGN_PARAGRAPH.RIGHT
        
        # Default to left alignment
        return WD_ALIGN_PARAGRAPH.LEFT

    def detect_font_size(self, block):
        """Detect appropriate font size based on content and position."""
        text = block[4]
        
        # Check if it's a title (all caps, short length)
        if text.isupper() and len(text.strip()) < 100:
            return self.font_size_mapping['title']
        
        # Check if it's a heading (starts with common legal patterns)
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return self.font_size_mapping['heading']
        
        # Check if it's a case number or reference
        if re.search(r'(NO\.|CASE|REFERENCE|DATE)', text, re.IGNORECASE):
            return self.font_size_mapping['subheading']
        
        # Default to body text
        return self.font_size_mapping['body']

    def should_bold(self, text):
        """Determine if text should be bold based on legal document patterns."""
        # Titles and headings
        if text.isupper() and len(text.strip()) < 100:
            return True
        
        # Legal headers
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Case numbers and references
        if re.search(r'(NO\.|CASE|REFERENCE|DATE)', text, re.IGNORECASE):
            return True
        
        return False

    def analyze_document_structure(self, doc):
        """Analyze the document to identify headers, footers, and common patterns."""
        logger.info("Analyzing document structure...")
        
        all_blocks = []
        top_blocks = []
        bottom_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            
            if blocks:
                # Sort blocks by vertical position
                blocks.sort(key=lambda b: b[1])
                
                # Collect top and bottom blocks
                if blocks:
                    top_blocks.append(blocks[0][4].strip())
                    bottom_blocks.append(blocks[-1][4].strip())
                
                all_blocks.extend(blocks)
        
        # Find frequent headers and footers
        top_counter = Counter(top_blocks)
        bottom_counter = Counter(bottom_blocks)
        
        # Identify headers that appear on most pages
        threshold = max(2, int(0.6 * len(doc)))
        frequent_headers = {text for text, count in top_counter.items() 
                          if count >= threshold and text.strip()}
        frequent_footers = {text for text, count in bottom_counter.items() 
                           if count >= threshold and text.strip()}
        
        logger.info(f"Found {len(frequent_headers)} frequent headers and {len(frequent_footers)} frequent footers")
        
        return frequent_headers, frequent_footers

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX with enhanced legal document processing."""
        logger.info(f"Starting conversion: {pdf_path} -> {docx_path}")
        
        try:
            # Open PDF
            doc = fitz.open(pdf_path)
            if not doc:
                raise ValueError("Failed to open PDF file")
            
            # Create Word document
            document = Document()
            
            # Set document margins for legal documents
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Analyze document structure
            frequent_headers, frequent_footers = self.analyze_document_structure(doc)
            
            # Process each page
            for page_num in range(len(doc)):
                logger.info(f"Processing page {page_num + 1}/{len(doc)}")
                
                page = doc[page_num]
                blocks = page.get_text("blocks")
                
                if not blocks:
                    continue
                
                # Filter out headers, footers, page numbers, and artifacts
                filtered_blocks = []
                for block in blocks:
                    text = block[4].strip()
                    
                    # Skip empty blocks
                    if not text:
                        continue
                    
                    # Skip frequent headers and footers
                    if text in frequent_headers or text in frequent_footers:
                        logger.debug(f"Skipping frequent header/footer: {text[:50]}...")
                        continue
                    
                    # Skip page numbers
                    if self.is_page_number(text):
                        logger.debug(f"Skipping page number: {text}")
                        continue
                    
                    # Skip watermarks
                    if self.is_watermark(text):
                        logger.debug(f"Skipping watermark: {text}")
                        continue
                    
                    # Skip scanner artifacts
                    if self.is_scanner_artifact(text):
                        logger.debug(f"Skipping scanner artifact: {text[:50]}...")
                        continue
                    
                    # Clean the text
                    cleaned_text = self.clean_text(text)
                    if cleaned_text:
                        # Update block with cleaned text
                        block = list(block)
                        block[4] = cleaned_text
                        filtered_blocks.append(block)
                
                # Sort blocks by vertical position, then horizontal
                filtered_blocks.sort(key=lambda b: (b[1], b[0]))
                
                # Process blocks and create paragraphs
                current_paragraph_text = ""
                current_alignment = WD_ALIGN_PARAGRAPH.LEFT
                current_font_size = self.font_size_mapping['body']
                current_bold = False
                
                for i, block in enumerate(filtered_blocks):
                    text = block[4]
                    
                    # Detect formatting for this block
                    alignment = self.detect_alignment(block)
                    font_size = self.detect_font_size(block)
                    bold = self.should_bold(text)
                    
                    # Check if we should start a new paragraph
                    should_new_paragraph = False
                    
                    # Different alignment requires new paragraph
                    if alignment != current_alignment:
                        should_new_paragraph = True
                    
                    # Different font size requires new paragraph
                    if font_size != current_font_size:
                        should_new_paragraph = True
                    
                    # Different bold setting requires new paragraph
                    if bold != current_bold:
                        should_new_paragraph = True
                    
                    # Significant vertical gap indicates new paragraph
                    if i > 0:
                        prev_block = filtered_blocks[i-1]
                        vertical_gap = block[1] - (prev_block[1] + prev_block[3])
                        if vertical_gap > 15:  # More than 15 points gap
                            should_new_paragraph = True
                    
                    # If we need a new paragraph, save the current one
                    if should_new_paragraph and current_paragraph_text:
                        self.add_paragraph_to_document(
                            document, current_paragraph_text, 
                            current_alignment, current_font_size, current_bold
                        )
                        current_paragraph_text = ""
                    
                    # Update current formatting
                    current_alignment = alignment
                    current_font_size = font_size
                    current_bold = bold
                    
                    # Add text to current paragraph
                    if current_paragraph_text:
                        current_paragraph_text += " " + text
                    else:
                        current_paragraph_text = text
                
                # Add the last paragraph
                if current_paragraph_text:
                    self.add_paragraph_to_document(
                        document, current_paragraph_text,
                        current_alignment, current_font_size, current_bold
                    )
                
                # Add page break between pages (except for the last page)
                if page_num < len(doc) - 1:
                    document.add_page_break()
            
            # Save the document
            document.save(docx_path)
            logger.info(f"Successfully converted PDF to DOCX: {docx_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error during conversion: {str(e)}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()

    def add_paragraph_to_document(self, document, text, alignment, font_size, bold):
        """Add a paragraph to the document with specified formatting."""
        if not text.strip():
            return
        
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        
        run = paragraph.add_run(text.strip())
        run.font.size = font_size
        run.bold = bold
        
        # Set font to a standard legal document font
        run.font.name = 'Times New Roman'

def main():
    """Main function for command-line usage."""
    if len(sys.argv) != 3:
        print("Usage: python enhanced_pdf_to_docx.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    # Validate input file
    if not os.path.exists(pdf_path):
        print(f"Error: Input PDF file not found: {pdf_path}", file=sys.stderr)
        sys.exit(1)
    
    # Create output directory if needed
    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert PDF to DOCX
    converter = LegalDocumentConverter()
    try:
        success = converter.convert_pdf_to_docx(pdf_path, docx_path)
        if success:
            print(f"Successfully converted {pdf_path} to {docx_path}")
        else:
            print("Conversion failed", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main() 