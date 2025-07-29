import sys
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter

def is_page_number(text):
    patterns = [
        r'\bpage\s*no\.?\s*\d+',
        r'\bpage\s*\d+',
        r'\d+\s*of\s*\d+',
        r'\(\d+\s*of\s*\d+\)',
        r'\[CW-\d+/\d+\]',
        r'\(\d+\s*of\s*\d+\)'
    ]
    return any(re.search(p, text, re.IGNORECASE) for p in patterns)

def is_header_line(text):
    header_keywords = [
        'high court', 'supreme court', 'CR. APP', 'dt.', r'of \d{4}', 'patna high court'
    ]
    return any(re.search(k, text, re.IGNORECASE) for k in header_keywords)

def is_bracketed_case_number(text):
    return bool(re.search(r'\[CW-\d+/\d+\]|\(\d+\s*of\s*\d+\)', text))

def clean_xml_unsafe(text):
    """
    Remove all control characters except \t, \n, \r to ensure XML compatibility.
    """
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

def clean_text(text):
    """
    Clean and normalize text, removing excessive whitespace while preserving structure.
    """
    if not text:
        return ""
    
    # Remove control characters
    text = clean_xml_unsafe(text)
    
    # Normalize whitespace - replace multiple spaces/tabs with single space
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Normalize line breaks - replace multiple newlines with single newline
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def should_add_paragraph_break(current_text, next_text, current_block, next_block):
    """
    Determine if we should add a paragraph break between two text blocks.
    """
    if not current_text or not next_text:
        return False
    
    # If there's significant vertical gap between blocks, add paragraph break
    if next_block and current_block:
        vertical_gap = next_block[1] - (current_block[1] + current_block[3])
        if vertical_gap > 10:  # More than 10 points gap
            return True
    
    # If current text ends with sentence-ending punctuation and next starts with capital
    if (current_text.rstrip().endswith(('.', '!', '?', ':')) and 
        next_text.lstrip() and next_text.lstrip()[0].isupper()):
        return True
    
    # If next text looks like a new paragraph (indented or starts with common paragraph starters)
    if re.match(r'^\s*[A-Z][a-z]+\s', next_text.lstrip()):
        return True
    
    return False

def pdf_to_docx_no_header(pdf_path, docx_path, header_height=60):
    doc = fitz.open(pdf_path)
    document = Document()
    
    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    top_block_texts = []
    all_blocks_by_page = []
    
    # First pass: collect top block text from each page
    for page in doc:
        blocks = page.get_text("blocks")
        all_blocks_by_page.append(blocks)
        if blocks:
            # Find the top-most block (smallest y1)
            top_block = min(blocks, key=lambda b: b[1])
            top_block_texts.append(top_block[4].strip())
    
    # Find frequent top blocks (headers/watermarks)
    block_counter = Counter(top_block_texts)
    threshold = max(2, int(0.7 * len(doc)))
    frequent_headers = {text for text, count in block_counter.items() if count >= threshold and text}
    
    for page_num, blocks in enumerate(all_blocks_by_page):
        # Remove blocks that are headers, page numbers, bracketed case numbers, or frequent top blocks
        filtered_blocks = []
        for b in blocks:
            text = b[4].strip()
            if not text:
                continue
            if text in frequent_headers:
                continue
            if is_page_number(text):
                continue
            # Removed is_header_line(text) check - keeping header content
            if is_bracketed_case_number(text):
                continue
            filtered_blocks.append(b)
        
        # Sort by vertical position, then horizontal
        filtered_blocks.sort(key=lambda b: (b[1], b[0]))
        
        # Collect all text from this page
        page_text_parts = []
        
        for block in filtered_blocks:
            block_text = clean_text(block[4])
            if block_text:
                page_text_parts.append(block_text)
        
        # Combine all text with proper spacing
        if page_text_parts:
            # Join text parts with single newlines, then clean up
            combined_text = '\n'.join(page_text_parts)
            
            # Remove excessive newlines and normalize spacing
            combined_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', combined_text)
            combined_text = re.sub(r'[ \t]+', ' ', combined_text)
            
            # Split into paragraphs and add only non-empty ones
            paragraphs = combined_text.split('\n\n')
            
            for para in paragraphs:
                para = para.strip()
                if para and len(para) > 1:  # Only add non-empty paragraphs with meaningful content
                    # Split paragraph into lines and join them properly
                    lines = [line.strip() for line in para.split('\n') if line.strip()]
                    if lines:
                        # Join lines with spaces (not newlines) to avoid extra breaks
                        paragraph_text = ' '.join(lines)
                        
                        # Create paragraph with proper formatting
                        paragraph = document.add_paragraph(paragraph_text)
                        
                        # Apply formatting based on content
                        if re.match(r'^[A-Z\s]+$', paragraph_text) and len(paragraph_text) > 10:
                            # All caps text - likely a header
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(14)
                        elif re.match(r'^\d+\.', paragraph_text):
                            # Numbered list item
                            paragraph.style = document.styles['List Number']
        
        # Add page break between pages (except for the last page)
        if page_num < len(all_blocks_by_page) - 1:
            document.add_page_break()
    
    document.save(docx_path)
    doc.close()

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python pdf_to_docx_no_header.py input.pdf output.docx", file=sys.stderr)
        sys.exit(1)
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    pdf_to_docx_no_header(pdf_path, docx_path) 