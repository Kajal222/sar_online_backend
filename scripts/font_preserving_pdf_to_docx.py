#!/usr/bin/env python3
"""Minimal PDF to DOCX converter that attempts to preserve font styling.

Uses PyMuPDF to read text spans and python-docx to replicate them while
maintaining font name and size. This approach focuses on direct span
mapping without heavy heuristics for legal formatting.
"""

import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    doc = fitz.open(pdf_path)
    word = Document()

    for page in doc:
        width = page.rect.width
        data = page.get_text("dict")
        for block in data.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                paragraph = word.add_paragraph()
                bbox = line.get("bbox", [0, 0, 0, 0])
                center = (bbox[0] + bbox[2]) / 2
                if abs(center - width / 2) < width * 0.1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif center > width * 0.6:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for span in line.get("spans", []):
                    run = paragraph.add_run(span.get("text", ""))
                    size = span.get("size")
                    font = span.get("font")
                    if size:
                        run.font.size = Pt(size)
                    if font:
                        run.font.name = font

        word.add_page_break()

    if word.paragraphs and word.paragraphs[-1].text == "":
        word.paragraphs[-1]._element.getparent().remove(word.paragraphs[-1]._element)

    word.save(docx_path)
    doc.close()


def main():
    if len(sys.argv) != 3:
        print("Usage: font_preserving_pdf_to_docx.py <input.pdf> <output.docx>")
        sys.exit(1)
    convert_pdf_to_docx(sys.argv[1], sys.argv[2])


if __name__ == "__main__":
    main()
