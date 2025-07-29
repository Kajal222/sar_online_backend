#!/usr/bin/env python3
"""
Legal Document PDF to DOCX Converter
====================================

Specialized converter for legal documents (judgments, orders, court documents)
that preserves exact formatting, alignment, and document structure.
"""

import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from collections import Counter, defaultdict
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class LegalDocumentConverter:
    """Specialized converter for legal documents with exact formatting preservation."""
    
    def __init__(self):
        self.page_number_patterns = [
            r'\bpage\s*no\.?\s*\d+',
            r'\bpage\s*\d+',
            r'\d+\s*of\s*\d+',
            r'\(\d+\s*of\s*\d+\)',
            r'\[CW-\d+/\d+\]',
            r'^\d+$',  # Standalone page numbers
            r'^\s*\d+\s*$',  # Page numbers with whitespace
            r'Page\s*\d+',
            r'PAGE\s*\d+',
            r'page\s*\d+'
        ]
        
        self.watermark_patterns = [
            r'CONFIDENTIAL',
            r'DRAFT',
            r'COPY',
            r'SCANNED',
            r'DIGITAL\s*COPY',
            r'ELECTRONIC\s*COPY',
            r'ORIGINAL',
            r'CERTIFIED\s*COPY',
            r'TRUE\s*COPY',
            r'OFFICIAL\s*COPY'
        ]
        
        # Legal document specific patterns
        self.legal_headers = [
            r'IN\s*THE\s*SUPREME\s*COURT\s*OF\s*INDIA',
            r'IN\s*THE\s*HIGH\s*COURT\s*OF\s*[A-Z\s]+',
            r'CRIMINAL\s*APPEAL\s*NO\.',
            r'CIVIL\s*APPEAL\s*NO\.',
            r'WRIT\s*PETITION\s*NO\.',
            r'SPECIAL\s*LEAVE\s*PETITION\s*NO\.',
            r'CRIMINAL\s*APPELLATE\s*JURISDICTION',
            r'CIVIL\s*APPELLATE\s*JURISDICTION'
        ]
        
        # Font size mapping for legal documents
        self.font_size_mapping = {
            'title': Pt(16),
            'heading': Pt(14),
            'subheading': Pt(12),
            'body': Pt(11),
            'small': Pt(10),
            'footnote': Pt(9)
        }

    def is_page_number(self, text):
        """Check if text is a page number."""
        if not text or len(text.strip()) < 1:
            return False
        
        text_clean = text.strip()
        
        # Check against patterns
        for pattern in self.page_number_patterns:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True
        
        # Check if it's just a number (likely page number)
        if re.match(r'^\d+$', text_clean):
            return True
            
        return False

    def is_watermark(self, text):
        """Check if text is a watermark."""
        if not text:
            return False
        
        text_upper = text.upper().strip()
        
        for pattern in self.watermark_patterns:
            if re.search(pattern, text_upper):
                return True
        
        return False

    def clean_text(self, text):
        """Clean text while preserving legal document formatting."""
        if not text:
            return ""
        
        # Remove control characters except tabs and newlines
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize whitespace but preserve intentional spacing
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()

    def detect_alignment(self, block, page_width=595):
        """Detect text alignment based on block position and content."""
        block_width = block[2] - block[0]  # Block width
        block_center = (block[0] + block[2]) / 2
        page_center = page_width / 2
        
        # If block spans most of the page width, it's likely justified
        if block_width > page_width * 0.8:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Check if text is centered (common for titles and headers)
        if abs(block_center - page_center) < page_width * 0.1:  # Within 10% of center
            return WD_ALIGN_PARAGRAPH.CENTER
        
        # Check if text is right-aligned (common for dates, case numbers)
        if block_center > page_center + page_width * 0.2:
            return WD_ALIGN_PARAGRAPH.RIGHT
        
        # Default to left alignment
        return WD_ALIGN_PARAGRAPH.LEFT

    def detect_font_size(self, text, block):
        """Detect appropriate font size based on content and position."""
        # Check if it's a title (all caps, short length, at top of page)
        if text.isupper() and len(text.strip()) < 100 and block[1] < 100:
            return self.font_size_mapping['title']
        
        # Check if it's a legal header
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return self.font_size_mapping['heading']
        
        # Check if it's a case number or reference
        if re.search(r'(NO\.|CASE|REFERENCE|DATE|APPEAL|PETITION)', text, re.IGNORECASE):
            return self.font_size_mapping['subheading']
        
        # Check if it's a party name (usually in all caps)
        if text.isupper() and len(text.strip()) < 50:
            return self.font_size_mapping['subheading']
        
        # Default to body text
        return self.font_size_mapping['body']

    def should_bold(self, text, block):
        """Determine if text should be bold based on legal document patterns."""
        # Titles and headings
        if text.isupper() and len(text.strip()) < 100:
            return True
        
        # Legal headers
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Case numbers and references
        if re.search(r'(NO\.|CASE|REFERENCE|DATE|APPEAL|PETITION)', text, re.IGNORECASE):
            return True
        
        # Party names (usually in all caps)
        if text.isupper() and len(text.strip()) < 50:
            return True
        
        # Document type indicators
        if text.strip().upper() in ['JUDGMENT', 'ORDER', 'REPORTABLE']:
            return True
        
        return False

    def should_underline(self, text):
        """Determine if text should be underlined."""
        # Document type indicators
        if text.strip().upper() in ['JUDGMENT', 'ORDER', 'REPORTABLE']:
            return True
        
        # Case numbers
        if re.search(r'NO\.\s*_+\s*OF\s*\d+', text, re.IGNORECASE):
            return True
        
        return False

    def analyze_document_structure(self, doc):
        """Analyze the document to identify headers, footers, and common patterns."""
        logger.info("Analyzing document structure...")
        
        all_blocks = []
        top_blocks = []
        bottom_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            
            if blocks:
                # Sort blocks by vertical position
                blocks.sort(key=lambda b: b[1])
                
                # Collect top and bottom blocks
                if blocks:
                    top_blocks.append(blocks[0][4].strip())
                    bottom_blocks.append(blocks[-1][4].strip())
                
                all_blocks.extend(blocks)
        
        # Find frequent headers and footers
        top_counter = Counter(top_blocks)
        bottom_counter = Counter(bottom_blocks)
        
        # Identify headers that appear on most pages
        threshold = max(2, int(0.6 * len(doc)))
        frequent_headers = {text for text, count in top_counter.items() 
                          if count >= threshold and text.strip()}
        frequent_footers = {text for text, count in bottom_counter.items() 
                           if count >= threshold and text.strip()}
        
        logger.info(f"Found {len(frequent_headers)} frequent headers and {len(frequent_footers)} frequent footers")
        
        return frequent_headers, frequent_footers

    def process_legal_document_structure(self, blocks, page_width=595):
        """Process blocks to maintain legal document structure."""
        processed_blocks = []
        
        for block in blocks:
            text = block[4].strip()
            if not text:
                continue
            
            # Clean the text
            cleaned_text = self.clean_text(text)
            if not cleaned_text:
                continue
            
            # Update block with cleaned text
            block = list(block)
            block[4] = cleaned_text
            
            # Add metadata for processing
            metadata = {
                'alignment': self.detect_alignment(block, page_width),
                'font_size': self.detect_font_size(cleaned_text, block),
                'bold': self.should_bold(cleaned_text, block),
                'underline': self.should_underline(cleaned_text),
                'is_legal_header': any(re.search(pattern, cleaned_text, re.IGNORECASE) for pattern in self.legal_headers),
                'is_party_name': cleaned_text.isupper() and len(cleaned_text.strip()) < 50,
                'is_case_number': bool(re.search(r'(NO\.|CASE|APPEAL|PETITION)', cleaned_text, re.IGNORECASE))
            }
            
            # Create new block with metadata
            new_block = list(block)
            new_block.append(metadata)
            processed_blocks.append(new_block)
        
        return processed_blocks

    def create_legal_document_structure(self, document, blocks):
        """Create proper legal document structure with correct formatting."""
        if not blocks:
            return
        
        # Sort blocks by vertical position, then horizontal
        blocks.sort(key=lambda b: (b[1], b[0]))
        
        current_section = None
        current_paragraph = None
        
        for i, block in enumerate(blocks):
            text = block[4]
            metadata = block[5]
            
            # Determine if this should be a new paragraph
            should_new_paragraph = True
            
            # Check if this block should be part of the previous paragraph
            if i > 0:
                prev_block = blocks[i-1]
                prev_metadata = prev_block[5]
                
                # Same formatting and small gap = same paragraph
                if (metadata['alignment'] == prev_metadata['alignment'] and
                    metadata['font_size'] == prev_metadata['font_size'] and
                    metadata['bold'] == prev_metadata['bold'] and
                    metadata['underline'] == prev_metadata['underline']):
                    
                    vertical_gap = block[1] - (prev_block[1] + prev_block[3])
                    if vertical_gap < 20:  # Small gap
                        should_new_paragraph = False
            
            if should_new_paragraph:
                # Create new paragraph
                paragraph = document.add_paragraph()
                paragraph.alignment = metadata['alignment']
                
                # Add text with formatting
                run = paragraph.add_run(text)
                run.font.size = metadata['font_size']
                run.bold = metadata['bold']
                if metadata['underline']:
                    run.underline = True
                
                # Set font to Times New Roman for legal documents
                run.font.name = 'Times New Roman'
                
                current_paragraph = paragraph
            else:
                # Add to existing paragraph
                if current_paragraph:
                    # Add space before new text
                    current_paragraph.add_run(" ")
                    run = current_paragraph.add_run(text)
                    run.font.size = metadata['font_size']
                    run.bold = metadata['bold']
                    if metadata['underline']:
                        run.underline = True
                    run.font.name = 'Times New Roman'
            
            # Add extra spacing for legal headers and section breaks
            if metadata['is_legal_header'] or metadata['is_case_number']:
                document.add_paragraph()  # Add blank line

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX with exact legal document formatting preservation."""
        logger.info(f"Starting legal document conversion: {pdf_path} -> {docx_path}")
        
        try:
            # Open PDF
            doc = fitz.open(pdf_path)
            if not doc:
                raise ValueError("Failed to open PDF file")
            
            # Create Word document
            document = Document()
            
            # Set document margins for legal documents
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Analyze document structure
            frequent_headers, frequent_footers = self.analyze_document_structure(doc)
            
            # Process each page
            for page_num in range(len(doc)):
                logger.info(f"Processing page {page_num + 1}/{len(doc)}")
                
                page = doc[page_num]
                blocks = page.get_text("blocks")
                
                if not blocks:
                    continue
                
                # Get page dimensions
                page_width = page.rect.width
                
                # Filter out headers, footers, page numbers, and artifacts
                filtered_blocks = []
                for block in blocks:
                    text = block[4].strip()
                    
                    # Skip empty blocks
                    if not text:
                        continue
                    
                    # Skip frequent headers and footers
                    if text in frequent_headers or text in frequent_footers:
                        logger.debug(f"Skipping frequent header/footer: {text[:50]}...")
                        continue
                    
                    # Skip page numbers
                    if self.is_page_number(text):
                        logger.debug(f"Skipping page number: {text}")
                        continue
                    
                    # Skip watermarks
                    if self.is_watermark(text):
                        logger.debug(f"Skipping watermark: {text}")
                        continue
                    
                    filtered_blocks.append(block)
                
                # Process blocks for legal document structure
                processed_blocks = self.process_legal_document_structure(filtered_blocks, page_width)
                
                # Create document structure
                self.create_legal_document_structure(document, processed_blocks)
                
                # Add page break between pages (except for the last page)
                if page_num < len(doc) - 1:
                    document.add_page_break()
            
            # Save the document
            document.save(docx_path)
            logger.info(f"Successfully converted legal document to DOCX: {docx_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error during conversion: {str(e)}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()

def main():
    """Main function for command-line usage."""
    if len(sys.argv) != 3:
        print("Usage: python legal_document_converter.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    # Validate input file
    if not os.path.exists(pdf_path):
        print(f"Error: Input PDF file not found: {pdf_path}", file=sys.stderr)
        sys.exit(1)
    
    # Create output directory if needed
    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert PDF to DOCX
    converter = LegalDocumentConverter()
    try:
        success = converter.convert_pdf_to_docx(pdf_path, docx_path)
        if success:
            print(f"Successfully converted {pdf_path} to {docx_path}")
        else:
            print("Conversion failed", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main() 