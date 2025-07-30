#!/usr/bin/env python3
"""
True OCR PDF to DOCX Converter for Legal Documents
==================================================

This script uses true OCR to read text in the exact order it appears in the PDF,
preserving the original document structure and formatting.
"""

import sys
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
from collections import Counter, defaultdict
import json
import logging
import cv2
import numpy as np
import pytesseract
from PIL import Image
import io

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class TrueOCRLegalDocumentConverter:
    """True OCR converter for legal PDF documents with exact structure preservation."""
    
    def __init__(self):
        self.page_number_patterns = [
            r'\bpage\s*no\.?\s*\d+',
            r'\bpage\s*\d+',
            r'\d+\s*of\s*\d+',
            r'\(\d+\s*of\s*\d+\)',
            r'\[CW-\d+/\d+\]',
            r'\(\d+\s*of\s*\d+\)',
            r'^\d+$',  # Standalone page numbers
            r'^\s*\d+\s*$',  # Page numbers with whitespace
            r'Page\s*\d+',
            r'PAGE\s*\d+',
            r'page\s*\d+'
        ]
        
        self.watermark_patterns = [
            r'CONFIDENTIAL',
            r'DRAFT',
            r'COPY',
            r'SCANNED',
            r'DIGITAL\s*COPY',
            r'ELECTRONIC\s*COPY',
            r'ORIGINAL',
            r'CERTIFIED\s*COPY',
            r'TRUE\s*COPY',
            r'OFFICIAL\s*COPY'
        ]
        
        # Legal document specific patterns
        self.legal_headers = [
            r'IN\s*THE\s*HIGH\s*COURT\s*OF\s*[A-Z\s]+',
            r'IN\s*THE\s*SUPREME\s*COURT\s*OF\s*[A-Z\s]+',
            r'HIGH\s*COURT\s*OF\s*JUDICATURE',
            r'BEFORE\s*THE\s*HON\'BLE\s*[A-Z\s]+',
            r'CRIMINAL\s*APPEAL\s*NO\.',
            r'CIVIL\s*APPEAL\s*NO\.',
            r'WRIT\s*PETITION\s*NO\.',
            r'SPECIAL\s*LEAVE\s*PETITION\s*NO\.'
        ]
        
        # Numbered list patterns
        self.numbered_list_patterns = [
            r'^\d+\.\s*',  # 1. text
            r'^\d+\)\s*',  # 1) text
            r'^[a-z]\)\s*',  # a) text
            r'^[A-Z]\)\s*',  # A) text
            r'^[ivxlcdm]+\.\s*',  # i. ii. iii. etc.
            r'^[IVXLCDM]+\.\s*',  # I. II. III. etc.
        ]
        
        # Font size mapping for legal documents
        self.font_size_mapping = {
            'title': Pt(16),
            'heading': Pt(14),
            'subheading': Pt(12),
            'body': Pt(11),
            'small': Pt(10),
            'footnote': Pt(9)
        }

    def is_page_number(self, text):
        """Check if text is a page number."""
        if not text or len(text.strip()) < 1:
            return False
        
        text_clean = text.strip()
        
        # Check against patterns
        for pattern in self.page_number_patterns:
            if re.search(pattern, text_clean, re.IGNORECASE):
                return True
        
        # Check if it's just a number (likely page number)
        if re.match(r'^\d+$', text_clean):
            return True
            
        return False

    def is_watermark(self, text):
        """Check if text is a watermark."""
        if not text:
            return False
        
        text_upper = text.upper().strip()
        
        for pattern in self.watermark_patterns:
            if re.search(pattern, text_upper):
                return True
        
        return False

    def clean_text(self, text):
        """Clean text while preserving legal document formatting."""
        if not text:
            return ""
        
        # Remove control characters except tabs and newlines
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize whitespace but preserve intentional spacing
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        return text.strip()

    def extract_text_with_true_ocr(self, page):
        """Extract text with true OCR in reading order."""
        # Convert page to high-resolution image
        mat = page.get_pixmap(matrix=fitz.Matrix(3, 3))  # Very high resolution for better OCR
        img_data = mat.tobytes("png")
        
        # Convert to PIL Image
        img = Image.open(io.BytesIO(img_data))
        
        # Convert to OpenCV format
        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        
        # Use Tesseract OCR with reading order preservation
        custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
        ocr_data = pytesseract.image_to_data(img_cv, config=custom_config, output_type=pytesseract.Output.DICT)
        
        # Extract text blocks with positioning in reading order
        text_blocks = []
        for i in range(len(ocr_data['text'])):
            if ocr_data['conf'][i] > 30:  # Lower confidence threshold for better coverage
                text = ocr_data['text'][i].strip()
                if text:
                    x = ocr_data['left'][i]
                    y = ocr_data['top'][i]
                    width = ocr_data['width'][i]
                    height = ocr_data['height'][i]
                    conf = ocr_data['conf'][i]
                    
                    text_blocks.append({
                        'text': text,
                        'x': x,
                        'y': y,
                        'width': width,
                        'height': height,
                        'confidence': conf,
                        'right': x + width,
                        'bottom': y + height
                    })
        
        # Sort by reading order (top to bottom, left to right)
        text_blocks.sort(key=lambda b: (b['y'], b['x']))
        
        return text_blocks

    def detect_court_header(self, text):
        """Detect if text is a court header."""
        court_patterns = [
            r'HIGH\s*COURT\s*OF\s*JUDICATURE',
            r'IN\s*THE\s*HIGH\s*COURT',
            r'IN\s*THE\s*SUPREME\s*COURT',
            r'SUPREME\s*COURT\s*OF\s*INDIA'
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in court_patterns)

    def detect_case_number(self, text):
        """Detect if text is a case number."""
        case_patterns = [
            r'CIVIL\s*WRIT\s*PETITION\s*NO\.',
            r'CRIMINAL\s*APPEAL\s*NO\.',
            r'CIVIL\s*APPEAL\s*NO\.',
            r'WRIT\s*PETITION\s*NO\.',
            r'S\.B\.\s*CIVIL\s*WRIT\s*PETITION'
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in case_patterns)

    def detect_numbered_list_with_ocr(self, text_blocks):
        """Detect numbered lists using OCR positioning data."""
        numbered_items = []
        
        for i, block in enumerate(text_blocks):
            text = block['text']
            
            # Check if this looks like a numbered list item
            for pattern in self.numbered_list_patterns:
                if re.match(pattern, text):
                    # Find continuation lines
                    continuation_lines = []
                    current_y = block['y']
                    number_x = block['x']
                    
                    # Look for lines that are indented and continue the same item
                    for j in range(i + 1, len(text_blocks)):
                        next_block = text_blocks[j]
                        
                        # Check if this is a continuation line
                        if (next_block['y'] > current_y and 
                            next_block['y'] <= current_y + 40 and  # Within reasonable vertical distance
                            next_block['x'] > number_x + 20):  # Indented from the number
                            
                            continuation_lines.append(next_block)
                            current_y = next_block['y']
                        elif next_block['y'] > current_y + 40:
                            # Too far down, probably next item
                            break
                    
                    numbered_items.append({
                        'number': text,
                        'number_x': block['x'],
                        'number_y': block['y'],
                        'content_blocks': continuation_lines,
                        'indent_level': self.calculate_indent_level(block['x'])
                    })
                    break
        
        return numbered_items

    def calculate_indent_level(self, x_position):
        """Calculate indent level based on x position."""
        # Define indent levels based on x position
        if x_position < 100:
            return 0  # No indent
        elif x_position < 150:
            return 1  # Small indent
        elif x_position < 200:
            return 2  # Medium indent
        else:
            return 3  # Large indent

    def detect_alignment_with_ocr(self, text_blocks, page_width):
        """Detect alignment using OCR positioning data."""
        if not text_blocks:
            return WD_ALIGN_PARAGRAPH.LEFT
        
        # Calculate center of page
        page_center = page_width / 2
        
        # Group blocks by vertical position (same line)
        line_groups = {}
        for block in text_blocks:
            y_key = int(block['y'] / 10) * 10  # Group by 10-pixel intervals
            if y_key not in line_groups:
                line_groups[y_key] = []
            line_groups[y_key].append(block)
        
        # Analyze each line
        alignments = []
        for y_key, blocks in line_groups.items():
            if len(blocks) == 1:
                # Single block - check its position
                block = blocks[0]
                block_center = block['x'] + block['width'] / 2
                
                if abs(block_center - page_center) < 50:
                    alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
                elif block_center > page_center + 100:
                    alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            else:
                # Multiple blocks - check if they span the page
                min_x = min(b['x'] for b in blocks)
                max_x = max(b['right'] for b in blocks)
                span_width = max_x - min_x
                
                if span_width > page_width * 0.8:
                    alignments.append(WD_ALIGN_PARAGRAPH.JUSTIFY)
                else:
                    alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
        
        # Return most common alignment
        if alignments:
            return max(set(alignments), key=alignments.count)
        
        return WD_ALIGN_PARAGRAPH.LEFT

    def detect_font_size(self, text):
        """Detect appropriate font size based on content."""
        # Check if it's a title (all caps, short length)
        if text.isupper() and len(text.strip()) < 100:
            return self.font_size_mapping['title']
        
        # Check if it's a heading (starts with common legal patterns)
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return self.font_size_mapping['heading']
        
        # Check if it's a case number or reference
        if re.search(r'(NO\.|CASE|REFERENCE|DATE)', text, re.IGNORECASE):
            return self.font_size_mapping['subheading']
        
        # Default to body text
        return self.font_size_mapping['body']

    def should_bold(self, text):
        """Determine if text should be bold based on legal document patterns."""
        # Court headers
        if self.detect_court_header(text):
            return True
        
        # Case numbers
        if self.detect_case_number(text):
            return True
        
        # Titles and headings
        if text.isupper() and len(text.strip()) < 100:
            return True
        
        # Legal headers
        for pattern in self.legal_headers:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        # Case numbers and references
        if re.search(r'(NO\.|CASE|REFERENCE|DATE)', text, re.IGNORECASE):
            return True
        
        return False

    def is_versus_section(self, text):
        """Check if text is part of a versus section."""
        versus_patterns = [
            r'VERSUS',
            r'VS\.',
            r'v\.',
            r'\.\.\.\s*PETITIONER',
            r'\.\.\.\s*APPELLANT',
            r'\.\.\.\s*RESPONDENT'
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in versus_patterns)

    def analyze_document_structure(self, doc):
        """Analyze the document to identify headers, footers, and common patterns."""
        logger.info("Analyzing document structure...")
        
        all_blocks = []
        top_blocks = []
        bottom_blocks = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            
            if blocks:
                # Sort blocks by vertical position
                blocks.sort(key=lambda b: b[1])
                
                # Collect top and bottom blocks
                if blocks:
                    top_blocks.append(blocks[0][4].strip())
                    bottom_blocks.append(blocks[-1][4].strip())
                
                all_blocks.extend(blocks)
        
        # Find frequent headers and footers
        top_counter = Counter(top_blocks)
        bottom_counter = Counter(bottom_blocks)
        
        # Identify headers that appear on most pages
        threshold = max(2, int(0.6 * len(doc)))
        frequent_headers = {text for text, count in top_counter.items() 
                          if count >= threshold and text.strip()}
        frequent_footers = {text for text, count in bottom_counter.items() 
                           if count >= threshold and text.strip()}
        
        logger.info(f"Found {len(frequent_headers)} frequent headers and {len(frequent_footers)} frequent footers")
        
        return frequent_headers, frequent_footers

    def convert_pdf_to_docx(self, pdf_path, docx_path):
        """Convert PDF to DOCX with true OCR processing."""
        logger.info(f"Starting true OCR conversion: {pdf_path} -> {docx_path}")
        
        try:
            # Open PDF
            doc = fitz.open(pdf_path)
            if not doc:
                raise ValueError("Failed to open PDF file")
            
            # Create Word document
            document = Document()
            
            # Set document margins for legal documents
            sections = document.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Analyze document structure
            frequent_headers, frequent_footers = self.analyze_document_structure(doc)
            
            # Process each page
            for page_num in range(len(doc)):
                logger.info(f"Processing page {page_num + 1}/{len(doc)} with true OCR")
                
                page = doc[page_num]
                page_width = page.rect.width
                
                # Extract text with true OCR in reading order
                text_blocks = self.extract_text_with_true_ocr(page)
                
                if not text_blocks:
                    continue
                
                # Filter out headers, footers, page numbers, and artifacts
                filtered_blocks = []
                for block in text_blocks:
                    text = block['text'].strip()
                    
                    # Skip empty blocks
                    if not text:
                        continue
                    
                    # Skip frequent headers and footers
                    if text in frequent_headers or text in frequent_footers:
                        logger.debug(f"Skipping frequent header/footer: {text[:50]}...")
                        continue
                    
                    # Skip page numbers
                    if self.is_page_number(text):
                        logger.debug(f"Skipping page number: {text}")
                        continue
                    
                    # Skip watermarks
                    if self.is_watermark(text):
                        logger.debug(f"Skipping watermark: {text}")
                        continue
                    
                    # Clean the text
                    cleaned_text = self.clean_text(text)
                    if cleaned_text:
                        block['text'] = cleaned_text
                        filtered_blocks.append(block)
                
                # Process blocks in reading order (top to bottom, left to right)
                for block in filtered_blocks:
                    text = block['text']
                    
                    # Detect formatting
                    alignment = self.detect_alignment_with_ocr([block], page_width)
                    font_size = self.detect_font_size(text)
                    bold = self.should_bold(text)
                    
                    # Special handling for court headers
                    if self.detect_court_header(text):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER
                        font_size = self.font_size_mapping['title']
                        bold = True
                    
                    # Special handling for case numbers
                    if self.detect_case_number(text):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER
                        font_size = self.font_size_mapping['subheading']
                        bold = True
                    
                    # Special handling for versus sections
                    if self.is_versus_section(text):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add paragraph to document
                    self.add_paragraph_to_document(
                        document, text,
                        alignment, font_size, bold
                    )
                
                # Add page break between pages (except for the last page)
                if page_num < len(doc) - 1:
                    document.add_page_break()
            
            # Save the document
            document.save(docx_path)
            logger.info(f"Successfully converted PDF to DOCX: {docx_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error during conversion: {str(e)}")
            raise
        finally:
            if 'doc' in locals():
                doc.close()

    def add_paragraph_to_document(self, document, text, alignment, font_size, bold):
        """Add a paragraph to the document with specified formatting."""
        if not text.strip():
            return
        
        paragraph = document.add_paragraph()
        paragraph.alignment = alignment
        
        run = paragraph.add_run(text.strip())
        run.font.size = font_size
        run.bold = bold
        
        # Set font to a standard legal document font
        run.font.name = 'Times New Roman'

def main():
    """Main function for command-line usage."""
    if len(sys.argv) != 3:
        print("Usage: python true_ocr_pdf_to_docx.py <input.pdf> <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    
    # Validate input file
    if not os.path.exists(pdf_path):
        print(f"Error: Input PDF file not found: {pdf_path}", file=sys.stderr)
        sys.exit(1)
    
    # Create output directory if needed
    output_dir = os.path.dirname(docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Convert PDF to DOCX
    converter = TrueOCRLegalDocumentConverter()
    try:
        success = converter.convert_pdf_to_docx(pdf_path, docx_path)
        if success:
            print(f"Successfully converted {pdf_path} to {docx_path}")
        else:
            print("Conversion failed", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main() 